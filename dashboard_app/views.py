from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib import messages
from django.contrib.auth.models import User
from django.core.files.storage import default_storage
from django.core.files import File
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.contrib import messages
from django.conf import settings
from django.urls import reverse
from django.db.models import Count
from django.db.models.functions import TruncMonth
from .models import Location, Item, MediaModel
import os
import requests
import boto3
from slugify import slugify
import psutil
import json
import zipfile
from datetime import datetime
import cv2
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime
from fastkml import kml
from shapely.geometry import Point, LineString, Polygon
from pyproj import Transformer
import xml.etree.ElementTree as ET

#MOBILE VERSION

def map_mobile(request):
    return render(request, 'dashboard_app/mobile/project_map.html')


def save_location_mobile(request):
    if request.method == 'POST':
        geojson = request.POST.get('geojson')
        name = request.POST.get('name')
        description = request.POST.get('description')

        # Vérifier les données reçues
        if not geojson or not name or not description:
            messages.warning(request, 'Assurez-vous de saisir correctement tous les champs.')
            return redirect(reverse('map_mobile'))

        # Vérifier si un projet avec le même nom existe déjà
        if Location.objects.filter(name=name).exists():
            messages.warning(request, 'Un projet avec ce nom existe déjà. Veuillez choisir un autre nom.')
            return redirect(reverse('map_mobile'))

        # Créer Le projet s'il n'existe pas encore
        location = Location.objects.create(
            name=name,
            description=description,
            geojson=geojson,
        )
        messages.success(request, 'Le projet a été enregistré avec succès!')
        return redirect(reverse('location_mobile', args=[location.id]))
    else:
        return redirect(reverse('map_mobile'))


def locations_list_mobile(request):
    month = request.GET.get('month')
    year = request.GET.get('year')
    archived = request.GET.get('archived')

    locations = Location.objects.all().order_by('-created_at')

    # Apply month and year filtering
    if month and not year:
        current_year = datetime.now().year
        locations = locations.filter(created_at__year=current_year, created_at__month=month)
    elif year:
        locations = locations.filter(created_at__year=year)
        if month:
            locations = locations.filter(created_at__month=month)

    # Apply archived filtering
    if archived == 'archived':
        locations = locations.filter(archived=True)
    elif archived == 'non_archived':
        locations = locations.filter(archived=False)

    paginator = Paginator(locations, 9)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    months = [{'name': datetime(2000, i, 1).strftime('%B'), 'value': i} for i in range(1, 13)]
    years = range(2020, datetime.now().year + 1)

    context = {
        'locations': page_obj,
        'page_obj': page_obj,
        'months': months,
        'years': years,
    }
    return render(request, 'dashboard_app/mobile/locations_list.html', context)


def location_detail_mobile(request, location_id):
    location = get_object_or_404(Location, pk=location_id)

    last_item = Item.objects.filter(location_id=location_id).order_by('-counter').first()
    last_number = last_item.counter + 1 if last_item else 1


    # Vérifier que les données GeoJSON sont analysées correctement
    try:
        location_geojson = json.loads(location.geojson)
    except json.JSONDecodeError as e:
        location_geojson = None
        print(f"Erreur lors du décodage du geojson : {e}")

    items = location.items.all()
    items_geojson = []
    for item in items:
        try:
            geojson = json.loads(item.geojson)
            properties = {
                "name": item.name,
                "description": item.description,
                "id": item.id,
                "media": []
            }

            media_files = item.media_files.all()
            for media in media_files:
                if media.file_url and media.file_url.startswith("https://"):
                    media_data = {
                        "id": media.id,
                        "file_url": media.file_url,  # URL du fichier
                        "file_type": "image" if media.file_url.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')) else "video",
                        "capture_angle": media.capture_angle,
                        "icon_url": media.icon_url,
                        "plot_image_url": media.plot_image_url if media.plot_image_url else None
                    }
                    properties["media"].append(media_data)  # Ajouter les données du média aux propriétés
                else:
                    print(f"URL manquante ou incorrecte pour le média {media.id} dans l'élément {item.id}")

            geojson['properties'] = properties
            items_geojson.append(geojson)
        except json.JSONDecodeError:
            print(f"Erreur lors du décodage du geojson pour l'élément {item.id}")
        except Exception as e:
            print(f"Erreur lors du traitement de l'élément {item.id} : {e}")

    items_geojson = json.dumps({"type": "FeatureCollection", "features": items_geojson})

    return render(request, 'dashboard_app/mobile/location_detail.html', {
        'location': location,
        'location_geojson': json.dumps(location_geojson),
        'items_geojson': items_geojson,
        'last_number': last_number,
    })

def add_item_mobile(request, location_id):
    """
    Cette vue permet d'ajouter un nouvel élément (Item) dans une location donnée.
    Elle gère le téléchargement des fichiers multimédias associés à cet élément.
    """
    location = get_object_or_404(Location, pk=location_id)

    if request.method == 'POST':
        geojson = request.POST.get('geojson')
        name = request.POST.get('name')
        description = request.POST.get('description')

        if not geojson or not name or not description:
            messages.warning(request, 'Veuillez vérifier que tous les champs sont correctement remplis.')
            return redirect(reverse('location-detail', args=[location.id]))

        # Créer et enregistrer le nouvel élément
        last_counter = Item.objects.filter(location=location).order_by('-counter').first()
        new_counter = (last_counter.counter + 1) if last_counter else 1

        # Créer l'élément et le sauvegarder dans la base de données
        item = Item.objects.create(
            name=name,
            description=description,
            location=location,
            geojson=geojson,
            counter=new_counter
        )

        # Traiter les fichiers multimédias
        media_files = request.FILES.getlist('media')
        if media_files:
            # Initialiser le client boto3 pour DigitalOcean Spaces
            s3 = boto3.client(
                's3',
                region_name='lon1',
                endpoint_url=settings.AWS_S3_ENDPOINT_URL,
                aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
            )

            for media in media_files:
                try:
                    # Créer un nom unique pour le fichier
                    unique_filename = f"uploads/{item.location.name}/{item.name}/{media.name}"
                    
                    # Télécharger le fichier sur DigitalOcean Spaces
                    s3.upload_fileobj(
                        media,
                        settings.AWS_STORAGE_BUCKET_NAME,
                        unique_filename,
                        ExtraArgs={'ACL': 'public-read'}
                    )

                    # Créer un lien direct pour le fichier
                    media_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"

                    # Sauvegarder le lien du fichier dans la base de données
                    MediaModel.objects.create(item=item, file_url=media_url)

                    print(f"Fichier téléchargé : {media.name}")

                except Exception as e:
                    print(f"Erreur lors du téléchargement du fichier {media.name} : {e}")
                    item.delete()  # Supprimer l'élément en cas d'échec
                    messages.warning(request, "Échec de téléchargement des fichiers multimédias, élément non enregistré.")
                    return redirect(reverse('location_mobile', args=[location.id]))

        messages.success(request, "L'élément et les fichiers multimédias ont été enregistrés avec succès!")
        return redirect(reverse('location_mobile', args=[location.id]))

    messages.warning(request, "Une erreur s'est produite lors de l'ajout de l'élément.")
    return redirect(reverse('location_mobile', args=[location.id]))


def update_item_mobile(request, item_id):
    """
    Cette vue permet de mettre à jour un élément existant (Item) et ses fichiers multimédias associés.
    Elle gère également la mise à jour des données GeoJSON et l'analyse des angles pour les fichiers images.
    """
    item = get_object_or_404(Item, id=item_id)

    if request.method == 'POST':
        # Traitement de la requête en JSON (pour les mises à jour via API)
        if request.content_type == 'application/json':
            try:
                data = json.loads(request.body)
                geojson = data.get('geojson')

                # Vérifier et mettre à jour les données GeoJSON
                if geojson:
                    item.geojson = json.dumps(geojson)
                    item.save()
                    return JsonResponse({'success': True, 'message': 'L\'élément a été mis à jour géographiquement avec succès!'})
                else:
                    return JsonResponse({'success': False, 'message': 'Veuillez vérifier que les données GeoJSON sont correctement entrées.'})
            except json.JSONDecodeError:
                return JsonResponse({'success': False, 'message': 'Erreur dans les données JSON.'})
            except Exception as e:
                return JsonResponse({'success': False, 'message': 'Une erreur inattendue s\'est produite: ' + str(e)})
        else:
            # Traitement des données via un formulaire HTML classique
            name = request.POST.get('name')
            description = request.POST.get('description')
            media_files = request.FILES.getlist('media')
            print(media_files)

            # Vérifier que tous les champs obligatoires sont remplis
            if not name or not description:
                messages.warning(request, 'Veuillez vérifier que tous les champs sont correctement remplis.')
                return redirect(reverse('location_mobile', args=[item.location.id]))

            # Mettre à jour les informations de l'élément
            item.name = name
            item.description = description
            item.save()

            # Initialiser le client boto3 pour DigitalOcean Spaces
            s3 = boto3.client(
                's3',
                region_name='lon1',
                endpoint_url=settings.AWS_S3_ENDPOINT_URL,
                aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
            )

            # Traiter les fichiers multimédias (images/vidéos)
            for media in media_files:
                try:
                    # Créer un nom de fichier unique pour l'espace de stockage
                    unique_filename = f"uploads/{item.location.name}/{item.name.replace(' ', '_')}/{media.name.replace(' ', '_')}"  # Remplace les espaces
                    # Télécharger le fichier vers DigitalOcean Spaces
                    s3.upload_fileobj(
                        media,
                        settings.AWS_STORAGE_BUCKET_NAME,
                        unique_filename,
                        ExtraArgs={'ACL': 'public-read'}  # rendre le fichier public
                    )

                    # Créer l'URL du fichier téléchargé
                    media_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"
                    print(media_url)

                    # Sauvegarder les détails du fichier multimédia dans la base de données
                    MediaModel.objects.create(
                        item=item,
                        file_url=media_url,
                        capture_angle=None,  # Remplissez avec vos données d'angle si nécessaire
                        plot_image_url=None   # Remplissez avec vos données d'image de tracé si nécessaire
                    )

                    print(f"Fichier téléchargé : {media.name}")

                except Exception as e:
                    print(f"Erreur lors du téléchargement du fichier {media.name} : {e}")

            # Notification de succès après mise à jour
            messages.success(request, "L'élément et les médias ont été mis à jour avec succès!")
            return redirect(reverse('location_mobile', args=[item.location.id]))

    return JsonResponse({'success': False, 'message': 'Requête non valide.'})


def update_media_mobile(request, media_id):
    media = get_object_or_404(MediaModel, id=media_id)
    new_file = request.FILES.get('media_file')  # الحصول على الملف الجديد من الطلب

    if new_file:
        print("new")
        # إعداد عميل boto3 لـ DigitalOcean Spaces
        s3 = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        # إذا كان يوجد ملف قديم، نحذفه من Spaces
        old_file_key = media.file_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "") if media.file_url else None

        # إنشاء اسم فريد للملف الجديد
        unique_filename = f"uploads/{media.item.location.name}/{media.item.name}/{new_file.name}"
        folder_prefix = f"uploads/{media.item.location.name}/{media.item.name}/"  # المجلد

        try:
            # رفع الملف الجديد إلى DigitalOcean Spaces
            s3.upload_fileobj(
                new_file,
                settings.AWS_STORAGE_BUCKET_NAME,
                unique_filename,
                ExtraArgs={'ACL': 'public-read'}
            )

            # إنشاء رابط جديد للملف المرفوع
            new_media_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"

            # تحديث قاعدة البيانات بالملف الجديد
            media.file_url = new_media_url
            media.save()

            # طباعة جميع الملفات في نفس المجلد
            response = s3.list_objects_v2(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Prefix=folder_prefix)
            if 'Contents' in response:
                print(f"Media files in the folder '{folder_prefix}':")
                for obj in response['Contents']:
                    print(f" - {obj['Key']} (Last Modified: {obj['LastModified']}, Size: {obj['Size']} bytes)")
            else:
                print(f"No files found in folder: {folder_prefix}")

            # محاولة حذف الملف القديم
            if old_file_key:
                print(f"Attempting to delete old file: {old_file_key}")

                # التحقق من وجود الملف قبل الحذف
                try:
                    response = s3.head_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=old_file_key)
                    print(f"Old file exists: {old_file_key}")
                    response = s3.delete_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=old_file_key)
                    print(f"Réponse de la suppression: {response}")
                    print(f"Ancien fichier supprimé: {old_file_key}")
                except s3.exceptions.ClientError as e:
                    # إذا كان الخطأ بسبب عدم وجود الملف
                    if e.response['Error']['Code'] == '404':
                        print(f"Le fichier n'existe pas: {old_file_key}")
                    else:
                        print(f"Erreur lors de la tentative de suppression: {e}")
                        messages.warning(request, f"Erreur lors de la suppression de l'ancien fichier: {e}")

            messages.success(request, "Le fichier média a été mis à jour avec succès!")
            return redirect(reverse('location_mobile', args=[media.item.location.id]))

        except Exception as e:
            messages.warning(request, f"Erreur lors du téléchargement du nouveau fichier: {e}")
            return redirect(reverse('location_mobile', args=[media.item.location.id]))

    else:
        messages.warning(request, 'Aucun nouveau fichier sélectionné pour la mise à jour.')

    return redirect(reverse('location_mobile', args=[media.item.location.id]))

def delete_media_mobile(request, media_id):
    """
    Cette vue permet de supprimer un fichier média spécifique de la base de données ainsi que 
    du stockage DigitalOcean Spaces.
    """
    if request.method == 'POST':
        media = get_object_or_404(MediaModel, pk=media_id)
        location_id = media.item.location.id

        # إعداد عميل boto3
        s3 = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        try:
            # استخراج مسار الملف من `file_url` فقط في حال كان موجودًا
            if media.file_url:
                # إزالة عنوان URL الأساسي للحصول على المسار داخل Spaces
                file_key = media.file_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "")
                s3.delete_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=file_key)

            # استخراج مسار الصورة البيانية من `plot_image_url` إذا كان موجودًا
            if media.plot_image_url:
                plot_key = media.plot_image_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "")
                s3.delete_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=plot_key)

            # حذف السجل من قاعدة البيانات
            media.delete()

            messages.success(request, 'Le fichier a été supprimé avec succès!')
        except Exception as e:
            messages.warning(request, f"Erreur lors de la suppression du fichier: {e}")

        return redirect(reverse('location_mobile', args=[location_id]))

    messages.warning(request, 'La méthode de commande est incorrecte.')
    return redirect('locations_mobile')




# WEB VERSION 
def dashboard(request):        
    # Calcul des statistiques clés
    total_locations = Location.objects.count()
    total_items = Item.objects.count()
    total_media = MediaModel.objects.count()

    # Statistiques mensuelles des projets et éléments
    monthly_projects = (
        Location.objects.filter(archived=False)
        .annotate(month=TruncMonth('created_at'))
        .values('month')
        .annotate(total=Count('id'))
        .order_by('month')
    )

    monthly_items = (
        Item.objects.annotate(month=TruncMonth('created_at'))
        .values('month')
        .annotate(total=Count('id'))
        .order_by('month')
    )

    # Préparation des données GeoJSON pour la carte
    locations = [
        {
            'type': 'Feature',
            'geometry': {
                'type': 'Point',
                'coordinates': ''
            },
            'properties': {
                'name': location.name,
                'description': location.description
            }
        }
        for location in Location.objects.all()
    ]

    # Extraire les données dans des listes pour le graphique
    months = [entry['month'].strftime('%b') for entry in monthly_projects]
    project_data = [entry['total'] for entry in monthly_projects]
    item_data = [entry['total'] for entry in monthly_items]
    
    all_messages = messages.get_messages(request)

    context = {
        'total_locations': total_locations,
        'total_items': total_items,
        'total_media': total_media,
        'locations': locations,
        'months': months,
        'project_data': project_data,
        'item_data': item_data,
        'messages': all_messages,
    }

    return render(request, 'dashboard_app/web/dashboard.html', context)

def map(request):
    return render(request, 'dashboard_app/web/project_map.html')

def locations_list(request):
    month = request.GET.get('month')
    year = request.GET.get('year')
    archived = request.GET.get('archived')

    locations = Location.objects.all().order_by('-created_at')

    # Apply month and year filtering
    if month and not year:
        current_year = datetime.now().year
        locations = locations.filter(created_at__year=current_year, created_at__month=month)
    elif year:
        locations = locations.filter(created_at__year=year)
        if month:
            locations = locations.filter(created_at__month=month)

    # Apply archived filtering
    if archived == 'archived':
        locations = locations.filter(archived=True)
    elif archived == 'non_archived':
        locations = locations.filter(archived=False)

    paginator = Paginator(locations, 9)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    months = [{'name': datetime(2000, i, 1).strftime('%B'), 'value': i} for i in range(1, 13)]
    years = range(2020, datetime.now().year + 1)

    context = {
        'locations': page_obj,
        'page_obj': page_obj,
        'months': months,
        'years': years,
    }
    return render(request, 'dashboard_app/web/locations_list.html', context)

def location_detail(request, location_id):
    location = get_object_or_404(Location, pk=location_id)
    last_item = Item.objects.filter(location_id=location_id).order_by('-counter').first()
    last_number = last_item.counter + 1 if last_item else 1

    # Vérifier que les données GeoJSON sont analysées correctement
    try:
        location_geojson = json.loads(location.geojson)
    except json.JSONDecodeError as e:
        location_geojson = None
        print(f"Erreur lors du décodage du geojson : {e}")

    items = location.items.all()
    items_geojson = []
    for item in items:
        try:
            geojson = json.loads(item.geojson)
            properties = {
                "name": item.name,
                "description": item.description,
                "id": item.id,
                "media": []
            }

            media_files = item.media_files.all()
            for media in media_files:
                if media.file_url and media.file_url.startswith("https://"):
                    media_data = {
                        "id": media.id,
                        "file_url": media.file_url,  # URL du fichier
                        "file_type": "image" if media.file_url.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')) else "video",
                        "capture_angle": media.capture_angle,
                        "icon_url": media.icon_url,
                        "plot_image_url": media.plot_image_url if media.plot_image_url else None
                    }
                    properties["media"].append(media_data)  # Ajouter les données du média aux propriétés
                else:
                    print(f"URL manquante ou incorrecte pour le média {media.id} dans l'élément {item.id}")

            geojson['properties'] = properties
            items_geojson.append(geojson)
        except json.JSONDecodeError:
            print(f"Erreur lors du décodage du geojson pour l'élément {item.id}")
        except Exception as e:
            print(f"Erreur lors du traitement de l'élément {item.id} : {e}")

    items_geojson = json.dumps({"type": "FeatureCollection", "features": items_geojson})

    return render(request, 'dashboard_app/web/location_detail.html', {
        'location': location,
        'location_geojson': json.dumps(location_geojson),
        'items_geojson': items_geojson,
        'last_number': last_number,
    })

def save_location(request):
    if request.method == 'POST':
        geojson = request.POST.get('geojson')
        name = request.POST.get('name')
        description = request.POST.get('description')

        # Vérifier les données reçues
        if not geojson or not name or not description:
            messages.warning(request, 'Assurez-vous de saisir correctement tous les champs.')
            return redirect(reverse('map'))

        # Vérifier si un projet avec le même nom existe déjà
        if Location.objects.filter(name=name).exists():
            messages.warning(request, 'Un projet avec ce nom existe déjà. Veuillez choisir un autre nom.')
            return redirect(reverse('map'))

        # Créer Le projet s'il n'existe pas encore
        location = Location.objects.create(
            name=name,
            description=description,
            geojson=geojson,
        )
        messages.success(request, 'Le projet a été enregistré avec succès!')
        return redirect(reverse('location-detail', args=[location.id]))
    else:
        return redirect(reverse('map'))

def update_location(request, location_id):
    location = get_object_or_404(Location, id=location_id)
    if request.method == 'POST':
        try:
            # Utilisation de request.POST pour obtenir les données du formulaire HTML
            geojson = request.POST.get('geojson')
            name = request.POST.get('name')
            description = request.POST.get('description')

            # Vérification de la validité des entrées
            if not geojson or not name or not description:
                messages.warning(request, "Veuillez vérifier que tous les champs sont correctement remplis.")
                return redirect(reverse('locations-list'))

            # Vérifier si un projet avec le même nom existe déjà
            if Location.objects.filter(name=name).exclude(id=location_id).exists():
                messages.warning(request, 'Un projet avec ce nom existe déjà. Veuillez choisir un autre nom.')
                return redirect(reverse('locations-list'))

            # Vérifier si un projet avec la même description existe déjà
            if Location.objects.filter(description=description).exclude(id=location_id).exists():
                messages.warning(request, 'Un projet avec cette description existe déjà. Veuillez choisir une autre description.')
                return redirect(reverse('locations-list'))

            # Mise à jour des données du lieu
            location.name = name
            location.description = description
            location.geojson = geojson
            location.save()
            
            messages.success(request, "Le projet a été mis à jour avec succès !")
            return redirect(reverse('locations-list'))
        except json.JSONDecodeError:
            # Gestion de l'erreur JSON au cas où les données du formulaire ne sont pas valides
            messages.warning(request, "Échec de la lecture des données JSON. Veuillez vérifier que les données sont correctement envoyées.")
            return redirect(reverse('locations-list'))
    return redirect(reverse('locations-list'))

def delete_location(request, location_id):
    location = get_object_or_404(Location, id=location_id)
    location.delete()
    messages.success(request, "Le projet a été supprimé avec succès.")
    return redirect('locations-list')

def toggle_archive_location(request, location_id):
    location = get_object_or_404(Location, pk=location_id)
    location.archived = not location.archived
    location.save()

    if location.archived:
        messages.success(request, "Le lieu a été archivé avec succès.")
    else:
        messages.success(request, "Le lieu a été désarchivé avec succès.")
    return redirect(reverse('locations-list'))

def analyze_image_angle(image_url):
    """
    Cette fonction analyse une image depuis une URL pour détecter des lignes et calculer l'angle moyen
    des lignes détectées. Elle utilise la détection des contours et des lignes pour estimer l'orientation de capture.
    """
    # Charger l'image à partir de l'URL
    response = requests.get(image_url)
    if response.status_code != 200:
        return {"error": "Erreur : Impossible de charger l'image."}
    
    image_data = np.asarray(bytearray(response.content), dtype="uint8")
    image = cv2.imdecode(image_data, cv2.IMREAD_COLOR)
    if image is None:
        return {"error": "Erreur : Image introuvable."}

    # Convertir l'image en niveaux de gris pour faciliter la détection des contours
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Appliquer la détection des contours de Canny
    edges = cv2.Canny(gray, 30, 200)

    # Détecter les lignes en utilisant la transformation de Hough
    lines = cv2.HoughLinesP(edges, 1, np.pi / 180, 50, minLineLength=50, maxLineGap=30)

    # Si aucune ligne n'est détectée, renvoyer une capture par défaut (face)
    if lines is None:
        return {
            "mean_angle": 0.0,
            "capture_direction": "Capture effectuée de face",
            "object_location": "L'objet est situé directement devant la caméra",
            "icon": "https://e7.pngegg.com/pngimages/792/964/png-clipart-computer-icons-up-arrow-angle-black.png"
        }

    # Extraire les angles des lignes détectées
    angles = []
    for line in lines:
        x1, y1, x2, y2 = line[0]
        angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))

        # Ignorer les lignes quasi-horizontales pour éviter les erreurs
        if abs(angle) > 5 and abs(angle) < 175:
            angles.append(angle)

    if not angles:
        return {
            "mean_angle": 0.0,
            "capture_direction": "Capture effectuée de face",
            "object_location": "L'objet est situé directement devant la caméra",
            "icon": "https://e7.pngegg.com/pngimages/792/964/png-clipart-computer-icons-up-arrow-angle-black.png"
        }

    # Calculer l'angle moyen
    mean_angle = np.median(angles) if len(angles) > 5 else np.mean(angles)
    angle_info = {"mean_angle": mean_angle}
    
    if -10 <= mean_angle <= 10:
        angle_info['capture_direction'] = "Capture effectuée de face"
        angle_info['object_location'] = "L'objet est situé directement devant la caméra"
        angle_info['icon'] = "https://e7.pngegg.com/pngimages/792/964/png-clipart-computer-icons-up-arrow-angle-black.png"
    elif mean_angle > 10:
        angle_info['capture_direction'] = "Capture effectuée du côté droit"
        angle_info['object_location'] = "L'objet est situé sur le côté droit de la caméra"
        angle_info['icon'] = "https://icons.iconarchive.com/icons/fa-team/fontawesome/256/FontAwesome-Angle-Right-icon.png"
    elif mean_angle < -10:
        angle_info['capture_direction'] = "Capture effectuée du côté gauche"
        angle_info['object_location'] = "L'objet est situé sur le côté gauche de la caméra"
        angle_info['icon'] = "https://cdn-icons-png.flaticon.com/512/8591/8591387.png"

    return angle_info

def create_angle_plot(angle, save_path):
    """
    Cette fonction crée un graphique représentant l'angle de capture détecté par la caméra.
    Un demi-cercle représentant 180° est dessiné, avec une ligne rouge marquant l'angle de capture.
    Le graphique est ensuite sauvegardé à l'emplacement spécifié.
    """
    # Vérifie si le répertoire de sauvegarde existe, sinon le créer
    directory = os.path.dirname(save_path)
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Initialiser le graphique avec une taille définie
    fig, ax = plt.subplots(figsize=(6, 6))

    # Dessiner un demi-cercle représentant 180 degrés (champ de vision de la caméra)
    theta = np.linspace(0, np.pi, 180)
    x = np.cos(theta)
    y = np.sin(theta)
    ax.plot(x, y, label="180° demi-cercle")

    # Calculer les coordonnées de la ligne pour marquer l'angle de capture
    radians = np.radians(angle)
    x_angle = [0, np.cos(radians)]
    y_angle = [0, np.sin(radians)]
    ax.plot(x_angle, y_angle, color='r', label=f"Angle de capture: {angle}°")

    # Ajuster l'apparence du graphique
    ax.set_aspect('equal')
    ax.set_title(f"Capture de la caméra: {angle:.2f}°")
    ax.set_xlim([-1.1, 1.1])
    ax.set_ylim([0, 1.1])
    ax.legend()

    # Sauvegarder le graphique à l'emplacement spécifié
    plt.savefig(save_path)
    plt.close(fig)


def add_item(request, location_id):
    """
    Cette vue permet d'ajouter un nouvel élément (Item) dans une location donnée.
    Elle gère le téléchargement des fichiers multimédias associés à cet élément.
    """
    location = get_object_or_404(Location, pk=location_id)

    if request.method == 'POST':
        geojson = request.POST.get('geojson')
        name = request.POST.get('name')
        description = request.POST.get('description')

        if not geojson or not name or not description:
            messages.warning(request, 'Veuillez vérifier que tous les champs sont correctement remplis.')
            return redirect(reverse('location-detail', args=[location.id]))

        # Créer et enregistrer le nouvel élément
        last_counter = Item.objects.filter(location=location).order_by('-counter').first()
        new_counter = (last_counter.counter + 1) if last_counter else 1

        # Créer l'élément dans la base de données
        item = Item.objects.create(
            name=name,
            description=description,
            location=location,
            geojson=geojson,
            counter=new_counter
        )

        # Traiter les fichiers multimédias
        media_files = request.FILES.getlist('media')
        if media_files:
            # Initialiser le client boto3
            s3 = boto3.client(
                's3',
                region_name='lon1',
                endpoint_url=settings.AWS_S3_ENDPOINT_URL,
                aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
            )

            for media in media_files:
                try:
                    # Créer un nom de fichier unique pour l'espace de stockage
                    unique_filename = f"uploads/{item.location.name}/{item.name.replace(' ', '_')}/{media.name.replace(' ', '_')}"
                    # Télécharger le fichier vers DigitalOcean Spaces
                    s3.upload_fileobj(
                        media,
                        settings.AWS_STORAGE_BUCKET_NAME,
                        unique_filename,
                        ExtraArgs={'ACL': 'public-read'}
                    )

                    # Créer l'URL du fichier téléchargé
                    media_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"
                    # Sauvegarder le lien dans la base de données
                    MediaModel.objects.create(item=item, file_url=media_url)

                    print(f"Fichier téléchargé : {media.name}")

                except Exception as e:
                    print(f"Erreur lors du téléchargement du fichier {media.name} : {e}")
                    # En cas d'échec du téléchargement, supprimer l'élément si nécessaire
                    item.delete()
                    messages.warning(request, "Échec du téléchargement des fichiers multimédias, élément non enregistré.")
                    return redirect(reverse('location-detail', args=[location.id]))

        messages.success(request, "L'élément et les fichiers multimédias ont été enregistrés avec succès!")
        return redirect(reverse('location-detail', args=[location.id]))

    messages.warning(request, "Une erreur s'est produite lors de l'ajout de l'élément.")
    return redirect(reverse('location-detail', args=[location.id]))

def update_item(request, item_id):
    """
    Cette vue permet de mettre à jour un élément existant (Item) et ses fichiers multimédias associés.
    Elle gère également la mise à jour des données GeoJSON et l'analyse des angles pour les fichiers images.
    """
    item = get_object_or_404(Item, id=item_id)

    if request.method == 'POST':
        # Traitement de la requête en JSON (pour les mises à jour via API)
        if request.content_type == 'application/json':
            try:
                data = json.loads(request.body)
                geojson = data.get('geojson')

                # Vérifier et mettre à jour les données GeoJSON
                if geojson:
                    item.geojson = json.dumps(geojson)
                    item.save()
                    return JsonResponse({'success': True, 'message': 'L\'élément a été mis à jour géographiquement avec succès!'})
                else:
                    return JsonResponse({'success': False, 'message': 'Veuillez vérifier que les données GeoJSON sont correctement entrées.'})
            except json.JSONDecodeError:
                return JsonResponse({'success': False, 'message': 'Erreur dans les données JSON.'})
            except Exception as e:
                return JsonResponse({'success': False, 'message': 'Une erreur inattendue s\'est produite: ' + str(e)})
        else:
            # Traitement des données via un formulaire HTML classique
            name = request.POST.get('name')
            description = request.POST.get('description')
            media_files = request.FILES.getlist('media')
            print(media_files)

            # Vérifier que tous les champs obligatoires sont remplis
            if not name or not description:
                messages.warning(request, 'Veuillez vérifier que tous les champs sont correctement remplis.')
                return redirect(reverse('location-detail', args=[item.location.id]))

            # Mettre à jour les informations de l'élément
            item.name = name
            item.description = description
            item.save()

            # Initialiser le client boto3 pour DigitalOcean Spaces
            s3 = boto3.client(
                's3',
                region_name='lon1',
                endpoint_url=settings.AWS_S3_ENDPOINT_URL,
                aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
            )

            # Traiter les fichiers multimédias (images/vidéos)
            for media in media_files:
                try:
                    # Créer un nom de fichier unique pour l'espace de stockage
                    unique_filename = f"uploads/{item.location.name}/{item.name.replace(' ', '_')}/{media.name.replace(' ', '_')}"  # Remplace les espaces
                    # Télécharger le fichier vers DigitalOcean Spaces
                    s3.upload_fileobj(
                        media,
                        settings.AWS_STORAGE_BUCKET_NAME,
                        unique_filename,
                        ExtraArgs={'ACL': 'public-read'}  # rendre le fichier public
                    )

                    # Créer l'URL du fichier téléchargé
                    media_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"
                    print(media_url)

                    # Sauvegarder les détails du fichier multimédia dans la base de données
                    MediaModel.objects.create(
                        item=item,
                        file_url=media_url,
                        capture_angle=None,  # Remplissez avec vos données d'angle si nécessaire
                        plot_image_url=None   # Remplissez avec vos données d'image de tracé si nécessaire
                    )

                    print(f"Fichier téléchargé : {media.name}")

                except Exception as e:
                    print(f"Erreur lors du téléchargement du fichier {media.name} : {e}")

            # Notification de succès après mise à jour
            messages.success(request, "L'élément et les médias ont été mis à jour avec succès!")
            return redirect(reverse('location-detail', args=[item.location.id]))

    return JsonResponse({'success': False, 'message': 'Requête non valide.'})


def delete_item(request, item_id):
    """Supprimer un élément spécifique."""
    if request.method == 'POST':
        item = get_object_or_404(Item, id=item_id)
        item.delete()
        return JsonResponse({'success': True, 'message': "L'élément a été supprimé avec succès!"})
    return JsonResponse({'error': 'Méthode de requête non valide.'}, status=400)

def update_media(request, media_id):
    media = get_object_or_404(MediaModel, id=media_id)
    new_file = request.FILES.get('media_file')  # الحصول على الملف الجديد من الطلب

    if new_file:
        # إعداد عميل boto3 لـ DigitalOcean Spaces
        s3 = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        # إذا كان يوجد ملف قديم، نحذفه من Spaces
        old_file_key = media.file_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "") if media.file_url else None

        # إنشاء اسم فريد للملف الجديد
        unique_filename = f"uploads/{media.item.location.name}/{media.item.name}/{new_file.name}"
        folder_prefix = f"uploads/{media.item.location.name}/{media.item.name}/"  # المجلد

        try:
            # رفع الملف الجديد إلى DigitalOcean Spaces
            s3.upload_fileobj(
                new_file,
                settings.AWS_STORAGE_BUCKET_NAME,
                unique_filename,
                ExtraArgs={'ACL': 'public-read'}
            )

            # إنشاء رابط جديد للملف المرفوع
            new_media_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"

            # تحديث قاعدة البيانات بالملف الجديد
            media.file_url = new_media_url
            media.save()

            # طباعة جميع الملفات في نفس المجلد
            response = s3.list_objects_v2(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Prefix=folder_prefix)
            if 'Contents' in response:
                print(f"Media files in the folder '{folder_prefix}':")
                for obj in response['Contents']:
                    print(f" - {obj['Key']} (Last Modified: {obj['LastModified']}, Size: {obj['Size']} bytes)")
            else:
                print(f"No files found in folder: {folder_prefix}")

            # محاولة حذف الملف القديم
            if old_file_key:
                try:
                    response = s3.delete_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=old_file_key)
                    print(f"Réponse de la suppression: {response}")
                    print(f"Ancien fichier supprimé: {old_file_key}")
                except Exception as e:
                    print(f"Erreur lors de la suppression de l'ancien fichier: {e}")
                    messages.warning(request, f"Erreur lors de la suppression de l'ancien fichier: {e}")

            messages.success(request, "Le fichier média a été mis à jour avec succès!")
            return redirect(reverse('location-detail', args=[media.item.location.id]))

        except Exception as e:
            messages.warning(request, f"Erreur lors du téléchargement du nouveau fichier: {e}")
            return redirect(reverse('location-detail', args=[media.item.location.id]))

    else:
        messages.warning(request, 'Aucun nouveau fichier sélectionné pour la mise à jour.')

    return redirect(reverse('location-detail', args=[media.item.location.id]))

def delete_media(request, media_id):
    """
    Cette vue permet de supprimer un fichier média spécifique de la base de données ainsi que 
    du stockage DigitalOcean Spaces.
    """
    if request.method == 'POST':
        media = get_object_or_404(MediaModel, pk=media_id)
        location_id = media.item.location.id

        # إعداد عميل boto3
        s3 = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        try:
            # استخراج مسار الملف من `file_url` فقط في حال كان موجودًا
            if media.file_url:
                # إزالة عنوان URL الأساسي للحصول على المسار داخل Spaces
                file_key = media.file_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "")
                s3.delete_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=file_key)

            # استخراج مسار الصورة البيانية من `plot_image_url` إذا كان موجودًا
            if media.plot_image_url:
                plot_key = media.plot_image_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "")
                s3.delete_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=plot_key)

            # حذف السجل من قاعدة البيانات
            media.delete()

            messages.success(request, 'Le fichier a été supprimé avec succès!')
        except Exception as e:
            messages.warning(request, f"Erreur lors de la suppression du fichier: {e}")

        return redirect(reverse('location-detail', args=[location_id]))

    messages.warning(request, 'La méthode de commande est incorrecte.')
    return redirect('locations-list')

def manage_media(request):
    """
    Cette vue permet de gérer tous les médias associés aux localisations et éléments. Elle calcule la taille totale 
    des fichiers médias et affiche les informations de stockage disponibles et utilisées.
    """
    # Récupérer toutes les localisations triées par la date de création
    locations = Location.objects.all().order_by('-created_at')

    # Récupérer tous les médias liés aux éléments en utilisant select_related pour minimiser les requêtes
    items = MediaModel.objects.select_related('item').all()

    media_data = {}  # Dictionnaire pour stocker les informations sur les fichiers médias par élément
    total_media_size = 0  # Variable pour stocker la taille totale de tous les fichiers médias

    # Parcourir chaque élément média et calculer la taille des fichiers
    for item in items:
        file_url = item.file_url  # Utiliser l'URL du fichier pour obtenir les informations
        file_size = 0  # Par défaut, on initialise la taille du fichier à 0

        try:
            # Récupérer les métadonnées du fichier directement à partir de DigitalOcean Spaces
            response = requests.head(file_url)
            file_size = int(response.headers.get('Content-Length', 0))  # Taille du fichier en octets
            total_media_size += file_size  # Ajouter la taille à la somme totale
        except Exception as e:
            print(f"Erreur lors de la récupération des métadonnées du fichier : {file_url}, Erreur : {e}")

        # Déterminer le type de fichier (image, vidéo, ou autre)
        media_type = 'image' if file_url.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')) else \
                     'video' if file_url.lower().endswith(('.mp4', '.avi', '.mov')) else 'document'

        # Ajouter les informations sur le fichier média dans le dictionnaire `media_data`
        media_data.setdefault(item.item.id, []).append({
            'id': item.id,
            'name': os.path.basename(file_url),  # Extraire le nom du fichier depuis l'URL
            'file_type': media_type,
            'size': file_size,
            'thumbnail_url': file_url  # Utiliser l'URL comme aperçu
        })

    # Utiliser psutil pour obtenir des informations sur l'utilisation du disque du serveur
    disk_usage = psutil.disk_usage('/')
    total_disk_size_gb = disk_usage.total / (1024 ** 3)  # Taille totale du disque en Go
    available_disk_size_gb = disk_usage.free / (1024 ** 3)  # Espace libre du disque en Go
    total_media_size_gb = total_media_size / (1024 ** 3)  # Taille totale des médias en Go

    # Calculer le pourcentage d'espace disque utilisé par les fichiers médias
    used_percentage = (total_media_size / disk_usage.total) * 100 if disk_usage.total > 0 else 0

    # Préparer le contexte à envoyer au template pour l'affichage
    context = {
        'locations': locations,  # Toutes les localisations
        'media_data': media_data,  # Dictionnaire contenant les informations sur les fichiers médias
        'total_media_size_gb': total_media_size_gb,  # Taille totale des fichiers médias en Go
        'total_disk_size_gb': total_disk_size_gb,  # Taille totale du disque en Go
        'available_disk_size_gb': available_disk_size_gb,  # Espace disque disponible en Go
        'used_percentage': used_percentage,  # Pourcentage d'espace utilisé par les fichiers médias
        'selected_item_id': request.GET.get('item_id')  # ID de l'élément sélectionné (facultatif)
    }

    # Renvoyer les données au template `manage_media.html` pour l'affichage
    return render(request, 'dashboard_app/web/manage_media.html', context)

def management_add_media(request, item_id):
    """
    Cette vue permet d'ajouter plusieurs fichiers multimédias (images/vidéos) à un élément spécifique.
    Elle analyse les images pour obtenir les angles de capture et génère des graphiques d'angles associés.
    """
    if request.method == 'POST' and request.FILES:
        item = get_object_or_404(Item, pk=item_id)
        files = request.FILES.getlist('mediaFiles')
        
        if not files:
            messages.warning(request, 'Aucun fichier sélectionné.')
            return redirect(reverse('manage-media') + f'?item_id={item_id}')

        # Initialiser boto3 pour DigitalOcean Spaces
        s3 = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        for media in files:
            try:
                # Créer un nom de fichier unique
                unique_filename = f"uploads/{item.location.name}/{item.name}/{media.name}"

                # Télécharger le fichier vers DigitalOcean Spaces
                s3.upload_fileobj(
                    media,  # fichier en mémoire
                    settings.AWS_STORAGE_BUCKET_NAME,  # nom du bucket
                    unique_filename,  # chemin et nom du fichier
                    ExtraArgs={'ACL': 'public-read'}  # rendre le fichier public
                )

                # Créer l'URL du fichier média
                media_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"

                # Analyser l'angle de capture pour les images
                capture_angle = None
                plot_url = None
                if media.content_type.startswith('image/'):
                    angle_info = analyze_image_angle(media_url)
                    if "error" in angle_info:
                        messages.warning(request, angle_info["error"])
                    else:
                        capture_angle = angle_info.get("mean_angle")

                        # Générer un graphique d'angle et le sauvegarder dans DigitalOcean Spaces
                        plot_filename = f"uploads/plots/item_{item.id}/plot_angle_{capture_angle:.2f}.png"
                        plot_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{plot_filename}"

                        # Générer le graphique d'angle et le télécharger
                        create_angle_plot(capture_angle, plot_filename)

                        # Télécharger le graphique vers DigitalOcean Spaces
                        with open(plot_filename, "rb") as plot_file:
                            s3.upload_fileobj(
                                plot_file,
                                settings.AWS_STORAGE_BUCKET_NAME,
                                plot_filename,
                                ExtraArgs={'ACL': 'public-read'}
                            )
                        print(f"Graphique d'angle téléchargé pour {media.name} : {plot_url}")

                # Sauvegarder les détails du média dans la base de données
                MediaModel.objects.create(
                    item=item, 
                    file_url=media_url,  # Utiliser file_url pour l'URL du fichier
                    capture_angle=capture_angle, 
                    plot_image_url=plot_url  # Utiliser plot_image_url pour l'URL du graphique
                )

                print(f"Fichier téléchargé : {media.name}")

            except Exception as e:
                print(f"Erreur lors du téléchargement du fichier {media.name} : {e}")
                messages.warning(request, f"Erreur lors du téléchargement du fichier {media.name}")

        messages.success(request, f'{len(files)} fichier(s) ont été téléchargé(s) avec succès!')
        return redirect(reverse('manage-media') + f'?item_id={item_id}')

    messages.warning(request, 'La méthode de requête est incorrecte ou aucun fichier n\'a été fourni.')
    return redirect(reverse('manage-media') + f'?item_id={item_id}')

def management_update_media(request, media_id):
    """
    Cette vue permet de mettre à jour un fichier média existant en remplaçant l'ancien fichier par un nouveau.
    Elle analyse les nouvelles images pour obtenir l'angle de capture et génère un graphique d'angle si applicable.
    """
    media = get_object_or_404(MediaModel, id=media_id)  # Récupérer l'instance du fichier média à partir de l'ID
    item_id = media.item.id  # Récupérer l'ID de l'élément associé
    new_file = request.FILES.get('media_file')  # Récupérer le nouveau fichier envoyé dans la requête

    if new_file:
        # Initialiser boto3 pour communiquer avec DigitalOcean Spaces
        s3 = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        # Créer un nom de fichier unique pour le nouveau média
        unique_filename = f"uploads/{media.item.location.name}/{media.item.name}/{new_file.name}"

        # Télécharger le nouveau fichier vers DigitalOcean Spaces
        s3.upload_fileobj(
            new_file,
            settings.AWS_STORAGE_BUCKET_NAME,
            unique_filename,
            ExtraArgs={'ACL': 'public-read'}  # Rendre le fichier public
        )

        # Créer l'URL du nouveau fichier
        new_file_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}"

        capture_angle = None
        plot_url = None

        # Si le nouveau fichier est une image, analyser l'angle de capture
        if new_file.content_type.startswith('image/'):
            angle_info = analyze_image_angle(new_file_url)  # Analyser l'angle de l'image
            if "error" in angle_info:
                messages.warning(request, angle_info["error"])  # Avertir l'utilisateur en cas d'erreur
            else:
                capture_angle = angle_info.get("mean_angle")  # Extraire l'angle moyen

                # Générer le graphique de l'angle et le sauvegarder dans DigitalOcean Spaces
                plot_filename = f"plots/item_{media.item.id}/plot_angle_{capture_angle:.2f}.png"
                plot_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{plot_filename}"

                # Créer et sauvegarder le graphique d'angle
                create_angle_plot(capture_angle, plot_filename)

        # Mettre à jour le modèle de média avec les nouveaux fichiers
        media.file_url = new_file_url
        media.capture_angle = capture_angle
        media.plot_image_url = plot_url
        media.save()  # Enregistrer les modifications dans la base de données

        messages.success(request, "Les médias et l'analyse d'angle ont été mis à jour avec succès!")
        return redirect(reverse('manage-media') + f'?item_id={item_id}')

    messages.warning(request, 'Aucun nouveau fichier sélectionné.')
    return redirect(reverse('manage-media') + f'?item_id={item_id}')

def management_delete_media(request, media_id):
    """
    Cette vue permet de supprimer un fichier média spécifique de la base de données et de DigitalOcean Spaces.
    """
    if request.method == 'POST':
        # Récupérer l'instance du fichier média à partir de l'ID
        media = get_object_or_404(MediaModel, pk=media_id)
        item_id = media.item.id  # Récupérer l'ID de l'élément associé

        # Initialiser boto3 pour communiquer avec DigitalOcean Spaces
        s3 = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        # Supprimer le fichier de DigitalOcean Spaces
        try:
            # Extraire le chemin du fichier en supprimant l'URL de base
            file_key = media.file_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "")
            s3.delete_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=file_key)  # Supprimer le fichier
        except Exception as e:
            messages.warning(request, f"Erreur lors de la suppression du fichier de DigitalOcean Spaces : {e}")

        # Supprimer le média de la base de données
        media.delete()

        messages.success(request, 'Le fichier a été supprimé avec succès!')
        return redirect(reverse('manage-media') + f'?item_id={item_id}')

    messages.warning(request, 'La méthode de requête est incorrecte.')
    return redirect(reverse('manage-media') + f'?item_id={item_id}')


def manage_items(request):
    """
    Cette vue permet de gérer et d'afficher tous les éléments (items) associés aux localisations.
    Les éléments sont regroupés par localisation et les informations sont organisées pour l'affichage.
    """
    locations = Location.objects.all().order_by('-created_at')

    item_data = {}
    for location in locations:
        for item in location.items.all():
            # Organiser les données des items par localisation
            item_data.setdefault(location.id, []).append({
                'id': item.id,
                'name': item.name,
                'description': item.description,
                'created_at': item.created_at.strftime('%Y-%m-%d'),
            })

    context = {
        'locations': locations,
        'item_data': item_data,
        'selected_item_id': request.GET.get('item_id')
    }
    return render(request, 'dashboard_app/web/manage_items.html', context)

def management_update_item(request, item_id):
    """
    Cette vue permet de mettre à jour les informations d'un élément (item) existant.
    Elle vérifie la validité des données entrées avant d'enregistrer les modifications.
    """
    item = get_object_or_404(Item, id=item_id)
    location_id = item.location.id

    if request.method == 'POST':
        # Récupérer les données du formulaire
        name = request.POST.get('item_name')
        description = request.POST.get('item_description')

        # Vérifier que les champs sont bien remplis
        if not name or not description:
            messages.warning(request, 'Veuillez vérifier que tous les champs sont correctement remplis.')
            return redirect(reverse('manage-items') + f'?location_id={location_id}')

        # Mettre à jour les informations de l'élément
        item.name = name
        item.description = description
        item.save()

        messages.success(request, "L'élément a été mis à jour avec succès!")
        return redirect(reverse('manage-items') + f'?location_id={location_id}')

    return JsonResponse({'success': False, 'message': 'Requête non valide.'})


def management_delete_item(request, item_id):
    """
    Cette vue permet de supprimer un élément (item) de la base de données.
    Elle vérifie la méthode de requête avant de procéder à la suppression.
    """
    item = get_object_or_404(Item, id=item_id)
    location_id = item.location.id

    if request.method == 'POST':
        item.delete()
        messages.success(request, "L'élément a été supprimé avec succès!")
        return redirect(reverse('manage-items') + f'?location_id={location_id}')

    messages.error(request, "Requête non valide.")
    return redirect(reverse('manage-items') + f'?location_id={location_id}')


def project_report(request):
    """
    Cette vue génère un rapport de projet basé sur les filtres de mois, année et statut d'archivage.
    Elle permet à l'utilisateur de filtrer les projets selon les critères sélectionnés.
    """
    month = request.GET.get('month')
    year = request.GET.get('year')
    archived = request.GET.get('archived')

    # Récupérer toutes les localisations, triées par date de création
    locations = Location.objects.all().order_by('-created_at')

    # Filtrage par mois et année
    if month and not year:
        # Si seul le mois est spécifié, on utilise l'année en cours par défaut
        current_year = datetime.now().year
        locations = locations.filter(created_at__year=current_year, created_at__month=month)
    elif year:
        # Si l'année est spécifiée, appliquer le filtre d'année
        locations = locations.filter(created_at__year=year)
        if month:
            # Si le mois est aussi spécifié, appliquer le filtre de mois
            locations = locations.filter(created_at__month=month)

    # Filtrage par statut d'archivage
    if archived == 'archived':
        locations = locations.filter(archived=True)
    elif archived == 'non_archived':
        locations = locations.filter(archived=False)

    # Pagination des résultats (9 localisations par page)
    paginator = Paginator(locations, 9)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Liste des mois et années pour les filtres dans le formulaire
    months = [{'name': datetime(2000, i, 1).strftime('%B'), 'value': i} for i in range(1, 13)]
    years = range(2020, datetime.now().year + 1)

    context = {
        'page_obj': page_obj,  # Objet de pagination pour gérer l'affichage des pages
        'months': months,      # Liste des mois pour le filtre
        'years': years,        # Liste des années pour le filtre
    }
    return render(request, 'dashboard_app/web/project_report.html', context)


def generate_location_word(request, location_id):
    """
    Cette vue génère un fichier ZIP contenant des documents Word pour chaque élément d'un emplacement donné.
    Les documents Word incluent les descriptions, les géométries sous forme de coordonnées, ainsi que les médias 
    (images et vidéos) liés à chaque élément.
    """
    # Récupérer les informations de l'emplacement spécifié
    location = get_object_or_404(Location, id=location_id)
    elements = Item.objects.filter(location=location)

    # Préparer un fichier zip pour contenir les documents et les médias
    today_date = datetime.now().strftime('%Y-%m-%d')
    zip_subdir = f"{today_date}_{location.name}"
    zip_filename = f"{zip_subdir}.zip"
    s = BytesIO()  # Stocker le fichier zip en mémoire
    zf = zipfile.ZipFile(s, "w")

    # Initialiser le client boto3 pour accéder à DigitalOcean Spaces
    s3 = boto3.client('s3',
                      region_name='lon1',
                      endpoint_url=settings.AWS_S3_ENDPOINT_URL,
                      aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                      aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY)

    # Créer un document Word pour chaque élément
    for element in elements:
        doc = Document()
        doc.add_heading(element.name, 0)  # Ajouter le titre de l'élément
        doc.add_paragraph(f"Description: {element.description}")  # Ajouter la description de l'élément

        # Traiter les données GeoJSON au format ArcGIS
        if element.geojson:
            try:
                geojson_data = json.loads(element.geojson)  # Charger les données GeoJSON
                geometry = geojson_data.get('geometry', {})

                # Gérer les types de géométries
                if 'paths' in geometry:  # Probablement un LineString
                    paths = geometry.get('paths', [])
                    doc.add_heading("Type de Géométrie: LineString", level=1)
                    if paths:
                        for index, path in enumerate(paths):
                            for point_index, point in enumerate(path):
                                doc.add_paragraph(f"  Point {point_index + 1}: Longitude: {point[0]}, Latitude: {point[1]}")
                    else:
                        doc.add_paragraph("Données de chemin non trouvées.")

                elif 'rings' in geometry: 
                    rings = geometry.get('rings', [])
                    doc.add_heading("Type de Géométrie: Polygon", level=1)
                    if rings:
                        for ring_index, ring in enumerate(rings):
                            doc.add_paragraph(f"Anneau {ring_index + 1}:")
                            for point_index, point in enumerate(ring):
                                if point_index == len(ring) - 1 and point == ring[0]:
                                    point_label = "Point 1 (Identique au Point de départ)"
                                else:
                                    point_label = f"Point {point_index + 1}"
                                doc.add_paragraph(f"  {point_label}: Longitude: {point[0]}, Latitude: {point[1]}")
                    else:
                        doc.add_paragraph("Données d'anneaux non trouvées pour le polygone.")

                elif 'x' in geometry and 'y' in geometry:  # Probablement un Point
                    x = geometry['x']
                    y = geometry['y']
                    doc.add_heading("Type de Géométrie: Point", level=1)
                    doc.add_paragraph(f"Coordonnées (Longitude, Latitude): {x}, {y}")

                else:
                    doc.add_paragraph("Type de géométrie non pris en charge ou données de géométrie non trouvées.")

            except (KeyError, json.JSONDecodeError) as e:
                print(f"Erreur lors de l'analyse des données ArcGIS: {e}")
                doc.add_paragraph("Erreur lors de l'analyse des données géographiques.")

        # Ajouter les images et vidéos au document Word depuis DigitalOcean Spaces
        for media in element.media_files.all():
            try:
                media_key = media.file_url.replace(f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/", "")  # Chemin du fichier dans Spaces
                media_url = media.file_url

                if media.file_url.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')):
                    # Télécharger l'image depuis Spaces
                    image_data = s3.get_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=media_key)['Body'].read()
                    image_filename = os.path.basename(media_key)

                    # Sauvegarder l'image temporairement pour l'ajouter au document Word
                    with open(image_filename, 'wb') as temp_image:
                        temp_image.write(image_data)
                        doc.add_picture(image_filename, width=Inches(2))  # Ajouter l'image au document
                        doc.add_paragraph(f"Image: {image_filename}")  # Ajouter le nom de l'image
                    os.remove(image_filename)  # Supprimer le fichier temporaire

                elif media.file_url.endswith(('.mp4', '.avi', '.mov')):
                    doc.add_paragraph(f"Vidéo: {os.path.basename(media_key)}")  # Ajouter le nom de la vidéo
                    doc.add_paragraph("Le fichier vidéo est inclus dans le fichier ZIP.")  # Mentionner que la vidéo est dans le zip

                # Ajouter le fichier média au fichier zip
                media_data = s3.get_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME, Key=media_key)['Body'].read()
                zf.writestr(f"{zip_subdir}/{element.name}/{os.path.basename(media_key)}", media_data)

            except Exception as e:
                print(f"Erreur lors de l'ajout du média: {e}")

        # Sauvegarder le document Word en mémoire (sans l'enregistrer sur le serveur)
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Ajouter le document Word au fichier zip avec un nom de fichier unique
        zf.writestr(f"{zip_subdir}/{element.name}_{element.id}.docx", doc_io.read())

    # Fermer le fichier zip après avoir ajouté tous les fichiers
    zf.close()

    # Préparer la réponse pour télécharger directement le fichier zip
    response = HttpResponse(s.getvalue(), content_type="application/x-zip-compressed")
    response['Content-Disposition'] = f'attachment; filename="{zip_filename}"'

    return response

def elements_rapport(request):
    """
    Cette vue génère un rapport des éléments en fonction des filtres sélectionnés (mois, année, statut d'archivage, et localisation).
    Elle permet de télécharger les informations des éléments sélectionnés sous forme d'un fichier ZIP contenant des documents Word et médias associés.
    """
    if request.method == "GET":
        # Requête GET initiale pour afficher le formulaire de sélection des items
        month = request.GET.get('month')
        year = request.GET.get('year')
        archived = request.GET.get('archived')
        location_id = request.GET.get('location_id')

        # Récupérer toutes les localisations et appliquer un filtre d'archivage si précisé
        locations = Location.objects.all().order_by('-created_at')
        if archived == 'archived':
            locations = locations.filter(archived=True)
        elif archived == 'non_archived':
            locations = locations.filter(archived=False)

        # Récupérer les items liés à une localisation spécifique
        items = Item.objects.filter(location_id=location_id).prefetch_related('media_files')

        # Filtrage des items par mois et année
        if month and not year:
            current_year = datetime.now().year
            items = items.filter(created_at__year=current_year, created_at__month=month)
        elif year:
            items = items.filter(created_at__year=year)
            if month:
                items = items.filter(created_at__month=month)

        paginator = Paginator(locations, 9)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        months = [{'name': datetime(2000, i, 1).strftime('%B'), 'value': i} for i in range(1, 13)]
        years = range(2020, datetime.now().year + 1)

        context = {
            'locations': locations,
            'page_obj': page_obj,
            'months': months,
            'years': years,
        }

        return render(request, 'dashboard_app/web/elements_report.html', context)

    elif request.method == "POST":
        # Récupérer les IDs des items sélectionnés dans la requête POST
        selected_item_ids = request.POST.getlist('selected_items')
        if not selected_item_ids:
            messages.warning(request, 'Aucun élément sélectionné.')
            return redirect(reverse('elements-rapport'))

        # Obtenir les éléments sélectionnés
        items = Item.objects.filter(id__in=selected_item_ids).prefetch_related('media_files')

        # Récupérer les informations de la localisation en utilisant l'ID de localisation
        location_id = request.POST.get('location_id')
        if not location_id:
            return HttpResponse("Aucune localisation spécifiée.", status=400)
        
        location = get_object_or_404(Location, id=location_id)

        # Préparer le fichier ZIP
        today_date = datetime.now().strftime('%Y-%m-%d')
        zip_subdir = f"{today_date}_elements_projet_{location.name}"
        zip_filename = f"{zip_subdir}.zip"
        s = BytesIO()  # Stocker le fichier ZIP en mémoire
        zf = zipfile.ZipFile(s, "w")

        # Créer un document Word pour chaque élément sélectionné
        for item in items:
            doc = Document()
            doc.add_heading(item.name, 0)  # Ajouter le titre de l'élément
            doc.add_paragraph(f"Description: {item.description}")  # Ajouter la description de l'élément

            # Traiter les données GeoJSON au format ArcGIS
            if item.geojson:
                try:
                    geojson_data = json.loads(item.geojson)  # Charger les données ArcGIS
                    geometry = geojson_data.get('geometry', {})

                    # Gérer les types de géométrie
                    if 'paths' in geometry:  # Possiblement un LineString
                        paths = geometry.get('paths', [])
                        doc.add_heading("Type de Géométrie: LineString", level=1)
                        for index, path in enumerate(paths):
                            doc.add_paragraph(f"Chemin {index + 1}:")
                            for point_index, point in enumerate(path):
                                doc.add_paragraph(f"  Point {point_index + 1}: Longitude: {point[0]}, Latitude: {point[1]}")

                    elif 'rings' in geometry: 
                        rings = geometry.get('rings', [])
                        doc.add_heading("Type de Géométrie: Polygon", level=1)
                        if rings:
                            for ring_index, ring in enumerate(rings):
                                doc.add_paragraph(f"Anneau {ring_index + 1}:")
                                for point_index, point in enumerate(ring):
                                    if point_index == len(ring) - 1 and point == ring[0]:
                                        point_label = "Point 1 (Identique au Point de départ)"
                                    else:
                                        point_label = f"Point {point_index + 1}"
                                    doc.add_paragraph(f"  {point_label}: Longitude: {point[0]}, Latitude: {point[1]}")
                        else:
                            doc.add_paragraph("Données d'anneaux non trouvées pour le polygone.")

                    elif 'x' in geometry and 'y' in geometry:  # Possiblement un Point
                        x = geometry['x']
                        y = geometry['y']
                        doc.add_heading("Type de Géométrie: Point", level=1)
                        doc.add_paragraph(f"Coordonnées: Longitude: {x}, Latitude: {y}")

                except (KeyError, json.JSONDecodeError) as e:
                    print(f"Erreur dans les données GeoJSON: {e}")
                    doc.add_paragraph("Erreur dans les données géographiques.")

            # Ajouter les images et vidéos au document Word
            for media in item.media_files.all():
                
                try:
                    media_url = media.file_url  # Assurez-vous d'utiliser file_url
                    
                    response = requests.get(media_url)
                    if response.status_code == 200:
                        media_data = BytesIO(response.content)
                        if media.file_url.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')):
                            doc.add_picture(media_data, width=Inches(2))
                            doc.add_paragraph(f"Image: {os.path.basename(media.file_url)}")
                        elif media.file_url.endswith(('.mp4', '.avi', '.mov')):
                            doc.add_paragraph(f"Vidéo: {os.path.basename(media.file_url)}")
                except Exception as e:
                    print(f"Erreur média: {e}")

            # Sauvegarder le document Word en mémoire
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            zf.writestr(f"{zip_subdir}/{item.name}_{item.id}.docx", doc_io.read())

            # Ajouter les fichiers médias (images et vidéos) au ZIP
            for media in item.media_files.all():
                try:
                    media_url = media.file_url
                    response = requests.get(media_url)
                    if response.status_code == 200:
                        zf.writestr(f"{zip_subdir}/{item.name}/{os.path.basename(media.file_url)}", response.content)
                except Exception as e:
                    print(f"Erreur fichier ZIP: {e}")

        # Fermer le fichier ZIP après l'ajout de tous les fichiers
        zf.close()

        # Préparer la réponse pour le téléchargement direct du fichier ZIP
        response = HttpResponse(s.getvalue(), content_type="application/x-zip-compressed")
        response['Content-Disposition'] = f'attachment; filename={zip_filename}'

        return response

    return HttpResponse("Méthode de requête non valide.", status=400)



def export_kml_kmz(request):
    month = request.GET.get('month')
    year = request.GET.get('year')
    archived = request.GET.get('archived')

    locations = Location.objects.all().order_by('-created_at')

    # Apply month and year filtering
    if month and not year:
        current_year = datetime.now().year
        locations = locations.filter(created_at__year=current_year, created_at__month=month)
    elif year:
        locations = locations.filter(created_at__year=year)
        if month:
            locations = locations.filter(created_at__month=month)

    # Apply archived filtering
    if archived == 'archived':
        locations = locations.filter(archived=True)
    elif archived == 'non_archived':
        locations = locations.filter(archived=False)

    paginator = Paginator(locations, 9)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    months = [{'name': datetime(2000, i, 1).strftime('%B'), 'value': i} for i in range(1, 13)]
    years = range(2020, datetime.now().year + 1)

    context = {
        'page_obj': page_obj,
        'months': months,
        'years': years,
    }
    return render(request, 'dashboard_app/web/export_kml_kmz.html', context)


def import_kml_kmz(request):
    """
    Cette vue permet d'afficher les localisations filtrées par mois, année, et statut d'archivage,
    et fournit une interface pour importer des fichiers KML/KMZ.
    """
    month = request.GET.get('month')
    year = request.GET.get('year')
    archived = request.GET.get('archived')

    # Récupérer toutes les localisations et les trier par date de création
    locations = Location.objects.all().order_by('-created_at')

    # Appliquer le filtrage par mois et année
    if month and not year:
        # Utiliser l'année en cours si seule le mois est spécifié
        current_year = datetime.now().year
        locations = locations.filter(created_at__year=current_year, created_at__month=month)
    elif year:
        # Filtrer selon l'année spécifiée
        locations = locations.filter(created_at__year=year)
        if month:
            # Filtrer aussi selon le mois si précisé
            locations = locations.filter(created_at__month=month)

    # Appliquer le filtrage par statut d'archivage
    if archived == 'archived':
        locations = locations.filter(archived=True)
    elif archived == 'non_archived':
        locations = locations.filter(archived=False)

    # Pagination des localisations (9 localisations par page)
    paginator = Paginator(locations, 9)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Générer les options de mois et d'années pour les filtres dans l'interface
    months = [{'name': datetime(2000, i, 1).strftime('%B'), 'value': i} for i in range(1, 13)]
    years = range(2020, datetime.now().year + 1)

    context = {
        'page_obj': page_obj,  # Objet de pagination pour les localisations
        'months': months,      # Liste des mois pour le filtre
        'years': years,        # Liste des années pour le filtre
    }
    return render(request, 'dashboard_app/web/import_kml_kmz.html', context)

def upload_kml_kmz(request, location_id):
    """
    Cette vue permet de télécharger et traiter un fichier KML ou KMZ pour une localisation spécifique.
    Le fichier est analysé et importé dans le système pour la localisation donnée.
    """
    location = get_object_or_404(Location, id=location_id)
    
    if request.method == 'POST' and 'file' in request.FILES:
        # Récupérer le fichier téléchargé et vérifier l'extension
        uploaded_file = request.FILES['file']
        file_extension = uploaded_file.name.split('.')[-1].lower()

        # Vérifier si le fichier est au format KML
        if file_extension == 'kml':
            try:
                # Traiter le fichier KML
                process_kml_file(uploaded_file, location, request)
                messages.success(request, 'Le fichier KML a été importé avec succès !')
                return redirect('import_kml_kmz')
            except Exception as e:
                messages.error(request, f"Erreur lors de l'importation du fichier KML: {e}")
                return redirect('import_kml_kmz')

        # Vérifier si le fichier est au format KMZ
        elif file_extension == 'kmz':
            try:
                # Traiter le fichier KMZ
                process_kmz_file(uploaded_file, location, request)
                messages.success(request, 'Le fichier KMZ a été importé avec succès !')
                return redirect('import_kml_kmz')
            except Exception as e:
                messages.error(request, f"Erreur lors de l'importation du fichier KMZ: {e}")
                return redirect('import_kml_kmz')

        # Si le fichier n'est ni KML ni KMZ, afficher une erreur
        else:
            messages.error(request, 'Le fichier téléchargé n\'est pas au format KML ou KMZ.')
            return redirect('import_kml_kmz')

    # Si aucun fichier n'est sélectionné, afficher un message d'erreur
    messages.error(request, 'Veuillez sélectionner un fichier à télécharger.')
    return redirect('import_kml_kmz')

def process_kml_file(kml_file, location, request, media_files=None):
    """
    Cette fonction traite un fichier KML pour extraire les éléments géographiques et les importe 
    en tant qu'éléments (items) liés à une localisation spécifique. Les données géométriques sont
    converties en format ArcGIS ESRI et les fichiers média sont stockés dans DigitalOcean Spaces.
    """
    transformer = Transformer.from_crs("EPSG:4326", "EPSG:3857", always_xy=True)
    s3_client = boto3.client(
        's3',
        region_name='lon1',
        endpoint_url=settings.AWS_S3_ENDPOINT_URL,
        aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
        aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
    )

    try:
        # Analyser le fichier KML
        tree = ET.parse(kml_file)
        root = tree.getroot()
        namespace = {"kml": "http://www.opengis.net/kml/2.2"}
        document = root.find("kml:Document", namespace)

        if document is not None:
            existing_items = Item.objects.filter(location=location)
            existing_item_names = [item.name for item in existing_items]

            duplicate_items = []
            new_items = []

            # Parcourir chaque Placemark pour extraire les données géographiques
            for placemark in document.findall("kml:Placemark", namespace):
                name = placemark.find("kml:name", namespace).text if placemark.find("kml:name", namespace) is not None else "Sans nom"
                description = placemark.find("kml:description", namespace).text if placemark.find("kml:description", namespace) is not None else "Sans description"

                if name in existing_item_names:
                    duplicate_items.append(name)
                    continue

                geometry = None
                esri_data = None
                media_file_url = None

                # Traiter les fichiers médias si fournis
                if media_files:
                    extended_data = placemark.find("kml:ExtendedData", namespace)
                    if extended_data is not None:
                        for data in extended_data.findall("kml:Data", namespace):
                            if data.get("name") == "media_file":
                                media_file_name = data.find("kml:value", namespace).text
                                media_file = media_files.get(media_file_name)

                                if media_file:
                                    # Utiliser boto3 pour télécharger le fichier vers DigitalOcean Spaces
                                    unique_filename = f'uploads/{location.name}/{name}/{media_file_name}'
                                    s3_client.upload_fileobj(
                                        media_file,
                                        settings.AWS_STORAGE_BUCKET_NAME,
                                        unique_filename,
                                        ExtraArgs={'ACL': 'public-read'}
                                    )
                                    media_file_url = f'https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}'

                # Extraire la géométrie du point
                if placemark.find("kml:Point", namespace) is not None:
                    coordinates = placemark.find("kml:Point/kml:coordinates", namespace).text.strip()
                    lon, lat = [float(value) for value in coordinates.split(",")[:2]]
                    x, y = transformer.transform(lon, lat)
                    geometry = {
                        "spatialReference": {"latestWkid": 3857, "wkid": 102100},
                        "x": x,
                        "y": y
                    }

                elif placemark.find("kml:LineString", namespace) is not None:
                    coordinates = placemark.find("kml:LineString/kml:coordinates", namespace).text.strip()
                    paths = [[transformer.transform(float(coord.split(",")[0]), float(coord.split(",")[1])) for coord in coordinates.split()]]
                    geometry = {
                        "spatialReference": {"latestWkid": 3857, "wkid": 102100},
                        "paths": paths
                    }

                elif placemark.find("kml:Polygon", namespace) is not None:
                    coordinates = placemark.find("kml:Polygon/kml:outerBoundaryIs/kml:LinearRing/kml:coordinates", namespace).text.strip()
                    rings = [[transformer.transform(float(coord.split(",")[0]), float(coord.split(",")[1])) for coord in coordinates.split()]]
                    geometry = {
                        "spatialReference": {"latestWkid": 3857, "wkid": 102100},
                        "rings": rings
                    }

                if geometry:
                    esri_data = {
                        "geometry": geometry,
                        "symbol": {
                            "type": "esriSLS" if "paths" in geometry else "esriSFS",
                            "color": [255, 165, 0],
                            "width": 2 if "paths" in geometry else None,
                            "style": "esriSLSSolid" if "paths" in geometry else "esriSFSSolid"
                        },
                        "attributes": {},
                        "popupTemplate": None
                    }

                    item = Item.objects.create(
                        name=name,
                        description=description,
                        location=location,
                        geojson=json.dumps(esri_data)
                    )
                    new_items.append(name)

                    if media_file_url:
                        MediaModel.objects.create(
                            item=item,
                            file_url=media_file_url,  # Assurez-vous d'utiliser `file_url` pour le lien complet
                            created_at=datetime.now()
                        )

            if duplicate_items:
                messages.warning(request, f"Les éléments suivants sont déjà ajoutés: {', '.join(duplicate_items)}")
            if new_items:
                messages.success(request, f"Les nouveaux éléments ont été ajoutés avec succès: {', '.join(new_items)}")

    except ET.ParseError as e:
        print(f"Erreur lors de l'analyse du fichier KML: {e}")
        messages.error(request, "Erreur lors de l'analyse du fichier KML.")

def process_kmz_file(kmz_file, location, request):
    """
    Cette fonction traite un fichier KMZ en extrayant son contenu, en particulier les fichiers KML 
    et les fichiers multimédias (images/vidéos). Les fichiers KML sont traités et les fichiers multimédias 
    sont sauvegardés en tant que fichiers associés à une localisation spécifique dans DigitalOcean Spaces.
    """
    try:
        print("Début du traitement du fichier KMZ")
        media_files = {}

        # Configurer le client S3 pour DigitalOcean Spaces
        s3_client = boto3.client(
            's3',
            region_name='lon1',
            endpoint_url=settings.AWS_S3_ENDPOINT_URL,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )

        # Décompresser le fichier KMZ
        with zipfile.ZipFile(kmz_file, 'r') as z:
            print("Fichier KMZ décompressé avec succès.")
            print("Liste des fichiers dans le KMZ:", z.namelist())

            # Parcourir tous les fichiers dans le fichier KMZ
            for file_name in z.namelist():
                print(f"Traitement du fichier: {file_name}")
                
                # Si le fichier est un fichier KML, le traiter
                if file_name.endswith('.kml'):
                    with z.open(file_name) as kml_file:
                        print(f"Traitement du fichier KML: {file_name}")
                        process_kml_file(kml_file, location, request, media_files)
                
                # Si le fichier est un fichier multimédia (image/vidéo), le sauvegarder dans DigitalOcean Spaces
                elif file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.mp4', '.avi', '.mov')):
                    try:
                        print(f"Tentative de sauvegarde du fichier multimédia: {file_name}")
                        item = location.items.first()  # Utiliser le premier élément lié pour les médias
                        if item:
                            unique_filename = f'uploads/{location.name}/{item.name}/{file_name}'
                            
                            # Télécharger le fichier multimédia dans DigitalOcean Spaces
                            with z.open(file_name) as media_file:
                                s3_client.upload_fileobj(
                                    media_file,
                                    settings.AWS_STORAGE_BUCKET_NAME,
                                    unique_filename,
                                    ExtraArgs={'ACL': 'public-read'}
                                )
                            
                            # Construire l'URL du fichier multimédia dans DigitalOcean Spaces
                            media_url = f'https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{unique_filename}'
                            
                            # Enregistrer l'URL du fichier dans le dictionnaire des fichiers multimédias
                            media_files[file_name] = media_url

                            # Sauvegarder le lien média dans la base de données
                            MediaModel.objects.create(
                                item=item,
                                file_url=media_url,  # Utiliser `file_url` لحفظ رابط DigitalOcean
                                created_at=datetime.now()
                            )

                    except Exception as e:
                        print(f"Erreur lors de la sauvegarde du fichier multimédia {file_name}: {e}")

    except Exception as e:
        print(f"Erreur lors du traitement du fichier KMZ: {e}")

def generate_location_kml(request, location_id):
    """
    Cette vue génère un fichier KML pour un emplacement spécifique,
    le sauvegarde dans DigitalOcean Spaces, et fournit un lien de téléchargement.
    """
    # Récupérer les informations de l'emplacement spécifié
    location = get_object_or_404(Location, id=location_id)
    elements = Item.objects.filter(location=location)

    # Créer l'élément racine du fichier KML
    kml = ET.Element('kml', xmlns="http://www.opengis.net/kml/2.2")
    document = ET.SubElement(kml, 'Document')
    ET.SubElement(document, 'name').text = location.name

    # Ajouter des styles pour les polygones, lignes, et points
    polygon_style = ET.SubElement(document, 'Style', id="polygonStyle")
    poly_linestyle = ET.SubElement(polygon_style, 'LineStyle')
    ET.SubElement(poly_linestyle, 'color').text = 'FF0099CC'
    ET.SubElement(poly_linestyle, 'width').text = '3'
    poly_fillstyle = ET.SubElement(polygon_style, 'PolyStyle')
    ET.SubElement(poly_fillstyle, 'color').text = '99330000'

    line_style = ET.SubElement(document, 'Style', id="lineStyle")
    line_linestyle = ET.SubElement(line_style, 'LineStyle')
    ET.SubElement(line_linestyle, 'color').text = 'FF0099CC'
    ET.SubElement(line_linestyle, 'width').text = '5'

    point_style = ET.SubElement(document, 'Style', id="pointStyle")
    point_iconstyle = ET.SubElement(point_style, 'IconStyle')
    ET.SubElement(point_iconstyle, 'color').text = 'FF0099CC'
    ET.SubElement(point_iconstyle, 'scale').text = '2'
    icon = ET.SubElement(point_iconstyle, 'Icon')
    ET.SubElement(icon, 'href').text = 'http://maps.google.com/mapfiles/kml/shapes/placemark_circle.png'

    # Créer un transformateur pour convertir les coordonnées à WGS84
    transformer = Transformer.from_crs("EPSG:3857", "EPSG:4326", always_xy=True)

    # Ajouter les éléments au fichier KML
    for element in elements:
        placemark = ET.SubElement(document, 'Placemark')
        ET.SubElement(placemark, 'name').text = element.name
        ET.SubElement(placemark, 'description').text = element.description

        if element.geojson:
            try:
                geojson_data = json.loads(element.geojson)
                geometry = geojson_data.get('geometry', {})

                if 'paths' in geometry:
                    # Ajouter une ligne au KML si la géométrie est de type LineString
                    line_string = ET.SubElement(placemark, 'LineString')
                    coordinates = ET.SubElement(line_string, 'coordinates')
                    path_coords = [
                        f"{transformer.transform(point[0], point[1])[0]},{transformer.transform(point[0], point[1])[1]},0"
                        for path in geometry.get('paths', []) for point in path
                    ]
                    coordinates.text = " ".join(path_coords)
                    ET.SubElement(placemark, 'styleUrl').text = '#lineStyle'

                elif 'rings' in geometry:
                    # Ajouter un polygone au KML si la géométrie est de type Polygon
                    polygon = ET.SubElement(placemark, 'Polygon')
                    outer_boundary = ET.SubElement(polygon, 'outerBoundaryIs')
                    linear_ring = ET.SubElement(outer_boundary, 'LinearRing')
                    coordinates = ET.SubElement(linear_ring, 'coordinates')
                    ring_coords = [
                        f"{transformer.transform(point[0], point[1])[0]},{transformer.transform(point[0], point[1])[1]},0"
                        for ring in geometry.get('rings', []) for point in ring
                    ]
                    coordinates.text = " ".join(ring_coords)
                    ET.SubElement(placemark, 'styleUrl').text = '#polygonStyle'

                elif 'x' in geometry and 'y' in geometry:
                    # Ajouter un point au KML si la géométrie est de type Point
                    point = ET.SubElement(placemark, 'Point')
                    coordinates = ET.SubElement(point, 'coordinates')
                    lon, lat = transformer.transform(geometry['x'], geometry['y'])
                    coordinates.text = f"{lon},{lat},0"
                    ET.SubElement(placemark, 'styleUrl').text = '#pointStyle'

            except (KeyError, json.JSONDecodeError) as e:
                print(f"Erreur lors de l'analyse des données ArcGIS: {e}")

    # Convertir l'arbre XML en chaîne de caractères
    kml_data = ET.tostring(kml, encoding='utf-8', method='xml')

    # Sauvegarder le fichier dans DigitalOcean Spaces
    s3 = boto3.client('s3', 
                      region_name='lon1', 
                      endpoint_url=settings.AWS_S3_ENDPOINT_URL, 
                      aws_access_key_id=settings.AWS_ACCESS_KEY_ID, 
                      aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY)
    kml_filename = f"{location.name}.kml"
    kml_file_path = f"exports/{kml_filename}"

    try:
        # Envoyer le fichier KML à DigitalOcean Spaces
        s3.put_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME,
                      Key=kml_file_path,
                      Body=kml_data,
                      ACL='public-read',
                      ContentType='application/vnd.google-earth.kml+xml')

        # Créer l'URL pour téléchargement
        kml_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{kml_file_path}"
        return redirect(kml_url)

    except Exception as e:
        print(f"Erreur lors de l'exportation du fichier KML vers DigitalOcean Spaces: {e}")
        return HttpResponse("Une erreur s'est produite lors de l'exportation du fichier KML.", status=500)


def generate_location_kmz(request, location_id):
    """
    Cette vue génère un fichier KMZ pour un emplacement spécifique et inclut les fichiers médias dans le KMZ.
    Le fichier KMZ est ensuite sauvegardé dans DigitalOcean Spaces.
    """
    # Récupérer l'emplacement et les éléments associés
    location = get_object_or_404(Location, id=location_id)
    elements = Item.objects.filter(location=location)

    # Créer l'élément racine du fichier KML
    kml = ET.Element('kml', xmlns="http://www.opengis.net/kml/2.2")
    document = ET.SubElement(kml, 'Document')
    ET.SubElement(document, 'name').text = location.name

    # Ajouter les styles pour les polygones, lignes, et points
    polygon_style = ET.SubElement(document, 'Style', id="polygonStyle")
    poly_linestyle = ET.SubElement(polygon_style, 'LineStyle')
    ET.SubElement(poly_linestyle, 'color').text = 'FF0099CC'
    ET.SubElement(poly_linestyle, 'width').text = '3'
    poly_fillstyle = ET.SubElement(polygon_style, 'PolyStyle')
    ET.SubElement(poly_fillstyle, 'color').text = '99330000'

    line_style = ET.SubElement(document, 'Style', id="lineStyle")
    line_linestyle = ET.SubElement(line_style, 'LineStyle')
    ET.SubElement(line_linestyle, 'color').text = 'FF0099CC'
    ET.SubElement(line_linestyle, 'width').text = '5'

    point_style = ET.SubElement(document, 'Style', id="pointStyle")
    point_iconstyle = ET.SubElement(point_style, 'IconStyle')
    ET.SubElement(point_iconstyle, 'color').text = 'FF0099CC'
    ET.SubElement(point_iconstyle, 'scale').text = '2'
    icon = ET.SubElement(point_iconstyle, 'Icon')
    ET.SubElement(icon, 'href').text = 'http://maps.google.com/mapfiles/kml/shapes/placemark_circle.png'

    # Configurer un transformateur pour les coordonnées en WGS84
    transformer = Transformer.from_crs("EPSG:3857", "EPSG:4326", always_xy=True)

    # Stocker les fichiers temporaires et multimédias
    kmz_memory = BytesIO()
    media_files = {}

    # Ajouter les éléments et leurs géométries au fichier KML
    for element in elements:
        placemark = ET.SubElement(document, 'Placemark')
        ET.SubElement(placemark, 'name').text = element.name
        description = ET.SubElement(placemark, 'description')
        description.text = element.description if element.description else ""

        if element.geojson:
            try:
                geojson_data = json.loads(element.geojson)
                geometry = geojson_data.get('geometry', {})

                if 'paths' in geometry:
                    line_string = ET.SubElement(placemark, 'LineString')
                    coordinates = ET.SubElement(line_string, 'coordinates')
                    path_coords = []
                    for path in geometry.get('paths', []):
                        for point in path:
                            lon, lat = transformer.transform(point[0], point[1])
                            path_coords.append(f"{lon},{lat},0")
                    coordinates.text = " ".join(path_coords)
                    ET.SubElement(placemark, 'styleUrl').text = '#lineStyle'

                elif 'rings' in geometry:
                    polygon = ET.SubElement(placemark, 'Polygon')
                    outer_boundary = ET.SubElement(polygon, 'outerBoundaryIs')
                    linear_ring = ET.SubElement(outer_boundary, 'LinearRing')
                    coordinates = ET.SubElement(linear_ring, 'coordinates')
                    ring_coords = []
                    for ring in geometry.get('rings', []):
                        for point in ring:
                            lon, lat = transformer.transform(point[0], point[1])
                            ring_coords.append(f"{lon},{lat},0")
                    coordinates.text = " ".join(ring_coords)
                    ET.SubElement(placemark, 'styleUrl').text = '#polygonStyle'

                elif 'x' in geometry and 'y' in geometry:
                    point = ET.SubElement(placemark, 'Point')
                    coordinates = ET.SubElement(point, 'coordinates')
                    lon, lat = transformer.transform(geometry['x'], geometry['y'])
                    coordinates.text = f"{lon},{lat},0"
                    ET.SubElement(placemark, 'styleUrl').text = '#pointStyle'

            except (KeyError, json.JSONDecodeError) as e:
                print(f"Erreur lors du traitement du GeoJSON: {e}")

        # Ajouter les fichiers médias associés dans le KMZ
        for media in element.media_files.all():
            media_basename = os.path.basename(media.file_url)
            try:
                response = requests.get(media.file_url)
                if response.status_code == 200:
                    media_files[media_basename] = response.content  # Lire le contenu du fichier média depuis l'URL
                    description.text += f'<br/><img src="{media_basename}" />'  # Inclure un lien dans la description
                else:
                    print(f"Erreur lors de la récupération du fichier {media.file_url}")
            except Exception as e:
                print(f"Erreur lors de la récupération du fichier média {media.file_url}: {e}")

    # Ajouter le fichier KML dans le KMZ
    with zipfile.ZipFile(kmz_memory, 'w', zipfile.ZIP_DEFLATED) as kmz:
        kmz.writestr("doc.kml", ET.tostring(kml, encoding='utf-8', method='xml'))
        for media_name, media_data in media_files.items():
            kmz.writestr(media_name, media_data)

    # Sauvegarder le KMZ dans DigitalOcean Spaces
    s3 = boto3.client('s3',
                      region_name='lon1',
                      endpoint_url=settings.AWS_S3_ENDPOINT_URL,
                      aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                      aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY)
    kmz_filename = f"{location.name}.kmz"
    kmz_file_path = f"exports/{kmz_filename}"

    try:
        s3.put_object(Bucket=settings.AWS_STORAGE_BUCKET_NAME,
                      Key=kmz_file_path,
                      Body=kmz_memory.getvalue(),
                      ACL='public-read',
                      ContentType='application/vnd.google-earth.kmz')

        # Créer le lien de téléchargement pour l'utilisateur
        kmz_url = f"https://{settings.AWS_S3_CUSTOM_DOMAIN}/africanmaps/{kmz_file_path}"
        return redirect(kmz_url)

    except Exception as e:
        print(f"Erreur lors de l'exportation du fichier KMZ vers DigitalOcean Spaces: {e}")
        return HttpResponse("Une erreur s'est produite lors de l'exportation du fichier KMZ.", status=500)


def user_list(request):
    """Liste tous les utilisateurs."""
    users = User.objects.all()  # Récupérer tous les utilisateurs

    # Gérer les actions du formulaire
    if request.method == 'POST':
        action = request.POST.get('action')
        user_id = request.POST.get('user_id')

        if action == 'add':  # Ajouter un nouvel utilisateur
            username = request.POST['username']
            first_name = request.POST['first_name']
            last_name = request.POST['last_name']
            email = request.POST['email']
            password1 = request.POST['password1']
            password2 = request.POST['password2']
            is_admin = request.POST.get('is_admin') == 'true'  # Vérifier si c'est un administrateur

            if password1 == password2:
                if User.objects.filter(username=username).exists():
                    messages.warning(request, "Le nom d'utilisateur existe déjà.")
                elif User.objects.filter(email=email).exists():
                    messages.warning(request, "L'email existe déjà.")
                else:
                    user = User.objects.create_user(
                        username=username,
                        first_name=first_name,
                        last_name=last_name,
                        email=email,
                        password=password1
                    )
                    if is_admin:
                        user.is_staff = True  # Rendre l'utilisateur administrateur
                        user.is_superuser = True  # Rendre l'utilisateur super administrateur
                    user.save()  # Enregistrer l'utilisateur
                    messages.success(request, "L'utilisateur a été ajouté avec succès.")
                    return redirect('user-list')
            else:
                messages.warning(request, "Les mots de passe ne correspondent pas.")
        
        elif action == 'delete':
            user = get_object_or_404(User, id=user_id)  # Récupérer l'utilisateur à supprimer
            if user.is_superuser and User.objects.filter(is_superuser=True).count() <= 1:
                messages.warning(request, "Vous ne pouvez pas supprimer le seul administrateur.")
            else:
                user.delete()  # Supprimer l'utilisateur
                messages.success(request, "L'utilisateur a été supprimé avec succès.")
        
        elif action == 'archive':
            user = get_object_or_404(User, id=user_id)  # Récupérer l'utilisateur à archiver
            user.is_active = False  # Archiver l'utilisateur
            user.save()
            messages.success(request, "L'utilisateur a été archivé avec succès.")
        
        elif action == 'activate':
            user = get_object_or_404(User, id=user_id)  # Récupérer l'utilisateur à activer
            user.is_active = True  # Activer l'utilisateur
            user.save()
            messages.success(request, "L'utilisateur a été activé avec succès.")

        elif action == 'edit':
            user = get_object_or_404(User, id=user_id)  # Récupérer l'utilisateur à modifier
            user.username = request.POST['username']
            user.email = request.POST['email']
            user.first_name = request.POST['first_name']
            user.last_name = request.POST['last_name']
            if request.POST['password']:
                user.set_password(request.POST['password'])  # Changer le mot de passe si fourni
            user.save()  # Enregistrer les modifications
            messages.success(request, "Les informations de l'utilisateur ont été mises à jour avec succès.")

    return render(request, 'dashboard_app/web/user_list.html', {'users': users})
